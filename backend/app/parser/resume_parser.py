"""Asynchronous resume document extraction and heuristic structured parsing."""

from __future__ import annotations

import asyncio
import io
import re
from collections.abc import Callable
from dataclasses import dataclass
from pathlib import Path

from app.models import (
    CertificationSchema,
    ContactSchema,
    EducationSchema,
    ExperienceSchema,
    ProjectSchema,
    ResumeSchema,
)


class ResumeParseError(ValueError):
    """Raised when a resume cannot be read or does not contain usable text."""


SECTION_ALIASES: dict[str, set[str]] = {
    "summary": {"summary", "professional summary", "profile", "objective"},
    "skills": {
        "skills",
        "technical skills",
        "core competencies",
        "competencies",
        "skills, tools & languages",
        "skills tools & languages",
    },
    "education": {"education", "academic background", "qualifications"},
    "experience": {
        "experience",
        "work experience",
        "employment history",
        "professional experience",
        "technical experience",
    },
    "projects": {
        "projects",
        "personal projects",
        "selected projects",
        "selected ai projects",
        "academic projects",
        "key projects",
    },
    "certifications": {"certifications", "certificates", "licenses"},
    "languages": {"languages", "language skills"},
}
EMAIL_PATTERN = re.compile(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b")
PHONE_PATTERN = re.compile(r"(?<!\w)(?:\+?\d[\d(). -]{7,}\d)(?!\w)")
URL_PATTERN = re.compile(r"https?://[^\s|]+|(?:www\.)[^\s|]+", re.IGNORECASE)
DATE_PATTERN = re.compile(
    r"(?:\b(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|"
    r"jul(?:y)?|aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)"
    r"\.?\s+)?\d{4}\b|\bpresent\b",
    re.IGNORECASE,
)
EMPLOYER_DATE_PATTERN = re.compile(r"^(?P<company>.+?)\s*\|\s*(?P<dates>.+?\d{4}.+)$")
ACTION_LINE_PATTERN = re.compile(
    r"^(?:architected|assessed|assisted|analyzed|applied|built|collaborated|collected|combined|"
    r"compared|confirmed|contributed|coordinated|developed|deployed|designed|diagnosed|documented|engineered|"
    r"established|executed|guided|implemented|integrated|leveraged|liaised|maintained|measured|"
    r"led|mentored|optimized|performed|prepared|provided|reviewed|simulated|supported|trained|transformed|"
    r"verified|visualized)\b",
    re.IGNORECASE,
)


@dataclass(frozen=True)
class ResumeDocument:
    """Uploaded document value object, decoupled from FastAPI's UploadFile."""

    filename: str
    content: bytes


class ResumeParser:
    """Reads supported resume files and maps their content into a ResumeSchema."""

    async def parse(self, document: ResumeDocument) -> ResumeSchema:
        """Extract and structure a PDF or DOCX resume without blocking the event loop."""
        text = await asyncio.to_thread(self._extract_text, document)
        normalized_text = self._normalize_text(text)
        if not normalized_text:
            raise ResumeParseError("The uploaded document does not contain extractable text")

        sections = self._detect_sections(normalized_text)
        return ResumeSchema(
            name=self._extract_name(normalized_text),
            contact=self._extract_contact(normalized_text),
            summary=self._first_section(sections, "summary"),
            skills=self._extract_skills(sections.get("skills", "")),
            education=self._extract_education(sections.get("education", "")),
            experience=self._extract_experience(sections.get("experience", "")),
            projects=self._extract_projects(sections.get("projects", "")),
            certifications=self._extract_certifications(sections.get("certifications", "")),
            sections=sections,
            raw_text=normalized_text,
            source_filename=document.filename,
        )

    @staticmethod
    def _extract_text(document: ResumeDocument) -> str:
        suffix = Path(document.filename).suffix.lower()
        extractors: dict[str, Callable[[bytes], str]] = {
            ".pdf": ResumeParser._extract_pdf_text,
            ".docx": ResumeParser._extract_docx_text,
        }
        extractor = extractors.get(suffix)
        if extractor is None:
            raise ResumeParseError("Only PDF and DOCX resumes are supported")
        try:
            return extractor(document.content)
        except ResumeParseError:
            raise
        except Exception as exc:
            raise ResumeParseError(f"Unable to read {suffix} resume") from exc

    @staticmethod
    def _extract_pdf_text(content: bytes) -> str:
        try:
            import fitz  # PyMuPDF

            with fitz.open(stream=content, filetype="pdf") as pdf:
                text = "\n".join(page.get_text() for page in pdf)
            if text.strip():
                return text
        except ImportError:
            pass

        try:
            import pdfplumber

            with pdfplumber.open(io.BytesIO(content)) as pdf:
                return "\n".join(page.extract_text() or "" for page in pdf.pages)
        except ImportError as exc:
            raise ResumeParseError("Install PyMuPDF or pdfplumber to parse PDF files") from exc

    @staticmethod
    def _extract_docx_text(content: bytes) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise ResumeParseError("Install python-docx to parse DOCX files") from exc
        document = Document(io.BytesIO(content))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        table_cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
        return "\n".join(paragraphs + table_cells)

    @staticmethod
    def _normalize_text(text: str) -> str:
        text = text.replace("\r", "\n").replace("\u2022", "•")
        lines = (re.sub(r"[ \t]+", " ", line).strip() for line in text.splitlines())
        return "\n".join(line for line in lines if line)

    def _detect_sections(self, text: str) -> dict[str, str]:
        indexed_headers: list[tuple[int, str]] = []
        lines = text.splitlines()
        for index, line in enumerate(lines):
            heading = line.lower().strip(" :.-")
            for canonical_name, aliases in SECTION_ALIASES.items():
                if heading in aliases:
                    indexed_headers.append((index, canonical_name))
                    break
        sections: dict[str, str] = {}
        for position, (start, name) in enumerate(indexed_headers):
            end = indexed_headers[position + 1][0] if position + 1 < len(indexed_headers) else len(lines)
            content = "\n".join(lines[start + 1 : end]).strip()
            if content:
                sections[name] = content
        return sections

    @staticmethod
    def _extract_name(text: str) -> str | None:
        lines = text.splitlines()
        header_lines = lines[:12]

        # PDF text extraction can place right-aligned dates before the visual header.
        # A candidate directly above a contact line is the most reliable name signal.
        for index, line in enumerate(header_lines[:-1]):
            next_line = header_lines[index + 1]
            if (EMAIL_PATTERN.search(next_line) or PHONE_PATTERN.search(next_line)) and ResumeParser._is_name_candidate(line):
                return line.title() if line.isupper() else line

        # Sidebar-based templates frequently emit contact and skills before the main
        # header. A name directly above an all-caps professional headline remains a
        # stable signal regardless of the extracted reading order.
        for index, line in enumerate(lines[:-1]):
            if ResumeParser._is_name_candidate(line) and ResumeParser._looks_like_headline(lines[index + 1]):
                return line.title() if line.isupper() else line

        for line in header_lines:
            if ResumeParser._is_name_candidate(line):
                return line.title() if line.isupper() else line
        return None

    @staticmethod
    def _is_name_candidate(line: str) -> bool:
        if EMAIL_PATTERN.search(line) or PHONE_PATTERN.search(line) or len(line) > 80:
            return False
        if any(character.isdigit() for character in line):
            return False
        words = line.split()
        return 2 <= len(words) <= 5 and all(word.replace("-", "").isalpha() for word in words)

    @staticmethod
    def _looks_like_headline(line: str) -> bool:
        normalized = line.replace("-", " ").upper()
        return (
            line.isupper()
            and len(line) <= 80
            and any(word in normalized for word in ("DEVELOPER", "ENGINEER", "DESIGNER", "ANALYST", "MANAGER", "CONSULTANT"))
        )

    @staticmethod
    def _extract_contact(text: str) -> ContactSchema:
        header = "\n".join(text.splitlines()[:8])
        urls = URL_PATTERN.findall(header)
        labeled_urls = re.findall(
            r"(?:github|gitlab|kaggle|hugging\s*face|portfolio|website)\s*:\s*([^\s|]+)",
            header,
            flags=re.IGNORECASE,
        )
        urls.extend(labeled_urls)
        normalized_urls = [url if url.startswith("http") else f"https://{url}" for url in urls]
        linkedin = next((url for url in normalized_urls if "linkedin.com" in url.lower()), None)
        portfolio = next((url for url in normalized_urls if url != linkedin), None)
        phone = next(
            (
                match.group(0)
                for match in PHONE_PATTERN.finditer(header)
                if not re.search(r"\d{4}\s*(?:-|–|—)\s*\d{4}", match.group(0))
            ),
            None,
        )
        return ContactSchema(
            email=(match.group(0) if (match := EMAIL_PATTERN.search(header)) else None),
            phone=phone,
            linkedin_url=linkedin,
            portfolio_url=portfolio,
        )

    @staticmethod
    def _bullet_items(lines: list[str]) -> list[str]:
        """Reconstruct bullet points whose text wrapped across PDF lines."""
        items: list[str] = []
        for line in lines:
            cleaned = line.strip()
            if not cleaned or re.fullmatch(r"\d{4}\s*(?:-|–|—)\s*(?:\d{4}|present)", cleaned, re.IGNORECASE):
                continue
            if cleaned.startswith(("•", "-", "*")):
                item = cleaned.lstrip("•-* ").strip()
                if item:
                    items.append(item)
            elif ACTION_LINE_PATTERN.match(cleaned):
                items.append(cleaned)
            elif items:
                items[-1] = f"{items[-1]} {cleaned}"
            else:
                items.append(cleaned)
        return items

    @staticmethod
    def _extract_skills(section: str) -> list[str]:
        if not section:
            return []
        skills: list[str] = []
        for line in section.splitlines():
            category_free = re.sub(r"^[A-Za-z/& ]{2,24}:\s*", "", line.strip())
            for candidate in re.split(r"[,;•](?![^()]*\))", category_free):
                skill = candidate.strip(" -")
                if skill.startswith("(") and skills:
                    skills[-1] = f"{skills[-1]} {skill}"
                elif 1 < len(skill) <= 80:
                    skills.append(skill)
        return list(dict.fromkeys(skills))

    def _extract_education(self, section: str) -> list[EducationSchema]:
        lines = [line.strip() for line in section.splitlines() if line.strip()]
        results: list[EducationSchema] = []
        for index, line in enumerate(lines):
            match = EMPLOYER_DATE_PATTERN.match(line)
            if not match:
                continue
            degree = lines[index - 1] if index else None
            start_date, end_date = self._split_dates(match.group("dates"))
            results.append(
                EducationSchema(
                    institution=match.group("company").strip(),
                    degree=degree,
                    start_date=start_date,
                    end_date=end_date,
                )
            )
        if results:
            return results
        for index, line in enumerate(lines):
            if re.search(r"\b(?:university|college|institute|school)\b", line, flags=re.IGNORECASE):
                return [EducationSchema(institution=line, degree=lines[index - 1] if index else None)]
        return [EducationSchema(institution=line) for line in lines[:1]]

    def _extract_experience(self, section: str) -> list[ExperienceSchema]:
        lines = [line.strip() for line in section.splitlines() if line.strip()]
        markers = [(index, EMPLOYER_DATE_PATTERN.match(line)) for index, line in enumerate(lines)]
        markers = [(index, match) for index, match in markers if match]
        results: list[ExperienceSchema] = []
        for marker_index, (line_index, match) in enumerate(markers):
            assert match is not None
            title = lines[line_index - 1] if line_index else "Not specified"
            next_marker_index = markers[marker_index + 1][0] if marker_index + 1 < len(markers) else len(lines)
            description_end = next_marker_index - 1 if marker_index + 1 < len(markers) else next_marker_index
            start_date, end_date = self._split_dates(match.group("dates"))
            results.append(
                ExperienceSchema(
                    company=match.group("company").strip(),
                    title=title,
                    start_date=start_date,
                    end_date=end_date,
                    description=self._bullet_items(lines[line_index + 1 : description_end]),
                )
            )
        if results:
            return results

        # Modern templates often list "job title" then "company" and render dates
        # in a separate column, so no company/date line exists in extracted text.
        role_company_pairs = [
            (index, index + 1)
            for index in range(len(lines) - 1)
            if self._looks_like_job_title(lines[index]) and self._looks_like_company(lines[index + 1])
        ]
        for position, (title_index, company_index) in enumerate(role_company_pairs):
            next_title_index = role_company_pairs[position + 1][0] if position + 1 < len(role_company_pairs) else len(lines)
            results.append(
                ExperienceSchema(
                    company=lines[company_index],
                    title=lines[title_index],
                    description=self._bullet_items(lines[company_index + 1 : next_title_index]),
                )
            )
        if results:
            return results

        # Some two-column PDFs extract dates separately from company and role. Retain
        # the experience evidence even when the date column cannot be paired safely.
        company_indexes = [
            index
            for index, line in enumerate(lines[:-1])
            if self._looks_like_company(line) and not lines[index + 1].startswith(("•", "-", "*"))
        ]
        for position, line_index in enumerate(company_indexes):
            next_index = company_indexes[position + 1] if position + 1 < len(company_indexes) else len(lines)
            title = lines[line_index + 1]
            description = self._bullet_items(lines[line_index + 2 : next_index])
            results.append(
                ExperienceSchema(
                    company=lines[line_index],
                    title=title,
                    description=description,
                )
            )
        return results

    @staticmethod
    def _looks_like_company(line: str) -> bool:
        normalized = line.lower()
        return len(line) >= 4 and any(
            token in normalized
            for token in ("co.", "ltd", "inc", "llc", "corp", "company", "innovations", "technologies", "solutions")
        )

    @staticmethod
    def _looks_like_job_title(line: str) -> bool:
        normalized = line.replace("-", " ").lower()
        return (
            not line.startswith(("•", "-", "*"))
            and len(line.split()) <= 8
            and any(word in normalized for word in ("developer", "engineer", "designer", "analyst", "manager", "consultant", "specialist", "intern"))
        )

    def _extract_projects(self, section: str) -> list[ProjectSchema]:
        projects: list[ProjectSchema] = []
        current_name: str | None = None
        current_lines: list[str] = []
        for line in section.splitlines():
            cleaned = line.strip()
            if not cleaned:
                continue
            if cleaned.startswith(("•", "-", "*")) or ACTION_LINE_PATTERN.match(cleaned):
                current_lines.append(cleaned)
            else:
                if cleaned.startswith("(") and current_name is not None:
                    current_lines.append(cleaned)
                elif current_name is not None and current_lines and not self._is_project_title(cleaned):
                    current_lines.append(cleaned)
                elif current_name is not None:
                    projects.append(ProjectSchema(name=current_name, description=self._bullet_items(current_lines)))
                    current_name, current_lines = cleaned, []
                else:
                    current_name, current_lines = cleaned, []
        if current_name is not None:
            projects.append(ProjectSchema(name=current_name, description=self._bullet_items(current_lines)))
        return projects

    @staticmethod
    def _is_project_title(line: str) -> bool:
        return bool(re.fullmatch(r"[A-Z][A-Za-z0-9/&() -]{2,}", line)) and not line.endswith(".")

    def _extract_certifications(self, section: str) -> list[CertificationSchema]:
        return [CertificationSchema(name=line.strip("•-* ")) for line in section.splitlines() if line.strip()]

    @staticmethod
    def _split_dates(value: str) -> tuple[str | None, str | None]:
        parts = re.split(r"\s*(?:–|—|-)\s*", value.strip(), maxsplit=1)
        return (parts[0], parts[1]) if len(parts) == 2 else (value.strip(), None)

    @staticmethod
    def _first_section(sections: dict[str, str], name: str) -> str | None:
        return sections.get(name) or None
